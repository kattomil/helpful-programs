from docx import Document
from docx.shared import Pt

from youtube_transcript_api import YouTubeTranscriptApi

import nltk
from nltk.corpus import wordnet as wn

video_id = input("Insert the video id: ")
subtitles = YouTubeTranscriptApi.get_transcript(video_id)

document = Document()

isverb = ["VB", "VBG", "VBD", "VBN", "VBP", "VBZ"]
verbs = []

for sub in subtitles:

    text = sub['text']
    
    for item in nltk.pos_tag(text.replace(".", "").replace(",", "").split()):
        if item[1] in isverb:
            verbs.append(item[0])

    p = document.add_paragraph(text)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 1.5
    
    p.style = document.styles['Normal']

writeverbs = "\nVerbs found: "
for verb in verbs:
    writeverbs += verb+", "
writeverbs += "\n"
    
p = document.add_paragraph(writeverbs)

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.line_spacing = 1.5
    
p.style = document.styles['Normal']

document.save('file.docx')

print("check the file")

# I_cEoK1mXms